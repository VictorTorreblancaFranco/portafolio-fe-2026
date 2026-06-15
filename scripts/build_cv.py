from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "src" / "assets" / "cv"
DOCX_PATH = OUT_DIR / "victor-torreblanca-cv.docx"

INK = "0B1220"
MUTED = "5B677A"
ACCENT = "0F766E"
LIGHT = "EAF7F4"
BORDER = "D9E2EC"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_hyperlink(paragraph, text, url, color=ACCENT):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    props.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title.upper())
    run.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def add_project(doc, title, meta, bullets, links=None):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(INK)
    title_run.font.size = Pt(10.5)
    p.add_run(f" | {meta}").font.color.rgb = RGBColor.from_string(MUTED)
    for bullet in bullets:
        add_bullet(doc, bullet)
    if links:
        link_p = doc.add_paragraph()
        link_p.paragraph_format.space_after = Pt(4)
        for idx, (label, url) in enumerate(links):
            if idx:
                link_p.add_run("  |  ")
            add_hyperlink(link_p, label, url)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(11.3)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(ACCENT)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(4)

    for style_name in ("List Bullet", "List Paragraph"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(9.6)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.left_indent = Inches(0.22)
        style.paragraph_format.first_line_indent = Inches(-0.12)


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    configure_styles(doc)

    header = doc.add_table(rows=1, cols=2)
    set_table_width(header, [Inches(4.55), Inches(2.75)])
    for row in header.rows:
        for cell in row.cells:
            set_cell_border(cell, "FFFFFF", "0")
            set_cell_margins(cell, top=20, bottom=20, start=20, end=20)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = header.cell(0, 0)
    name = left.paragraphs[0]
    name.paragraph_format.space_after = Pt(0)
    r = name.add_run("VICTOR RAUL TORREBLANCA FRANCO")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(INK)
    title = left.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    tr = title.add_run("Software Developer | Backend & Frontend Developer")
    tr.bold = True
    tr.font.size = Pt(11.5)
    tr.font.color.rgb = RGBColor.from_string(ACCENT)
    subtitle = left.add_paragraph("Estudiante de Análisis de Sistemas | Full Stack Developer Junior")
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    right = header.cell(0, 1)
    contact_lines = [
        ("WhatsApp: +51 930 574 584", "https://wa.me/51930574584"),
        ("GitHub", "https://github.com/VictorTorreblancaFranco"),
        ("LinkedIn", "https://www.linkedin.com/in/victor-raul-torreblanca-franco-687a04305"),
    ]
    for idx, (label, url) in enumerate(contact_lines):
        p = right.paragraphs[0] if idx == 0 else right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        add_hyperlink(p, label, url)

    intro = doc.add_table(rows=1, cols=1)
    set_table_width(intro, [Inches(7.34)])
    cell = intro.cell(0, 0)
    shade_cell(cell, LIGHT)
    set_cell_border(cell, "B7E4DA", "8")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.add_run(
        "Desarrollador junior en formación con enfoque en desarrollo web, APIs REST, "
        "microservicios, frontend Angular y despliegues con contenedores. Construyo experiencia "
        "real mediante proyectos académicos colaborativos, integrando backend, frontend, bases de "
        "datos y prácticas cloud native."
    )

    add_section_title(doc, "Tecnologías")
    tech_table = doc.add_table(rows=4, cols=2)
    set_table_width(tech_table, [Inches(1.55), Inches(5.79)])
    tech_rows = [
        ("Backend", "Java, Spring Boot, Spring WebFlux, Node.js, Python, APIs REST"),
        ("Frontend", "Angular, React, Flutter, HTML, CSS, SCSS"),
        ("Bases de datos", "PostgreSQL, MySQL, Oracle, MongoDB"),
        ("DevOps", "Docker, Kubernetes, Git, GitHub, GitLab, GitHub Actions, GitLab CI/CD, Kafka, Linux, YAML"),
    ]
    for row, (label, values) in zip(tech_table.rows, tech_rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=70, bottom=70, start=110, end=110)
        shade_cell(row.cells[0], "F3F8F7")
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(values)

    add_section_title(doc, "Proyectos destacados")
    add_project(
        doc,
        "EDUNOVA - Plataforma de Gestión Educativa",
        "Proyecto académico colaborativo",
        [
            "Ecosistema basado en microservicios para asistencia, usuarios, calificaciones, matrícula, bienestar, comunicación y gateway.",
            "Rol personal: responsable del microservicio de asistencia dentro del equipo.",
            "Stack: Java, Spring Boot, WebFlux, PostgreSQL, Docker, Angular y GitLab.",
        ],
        [("Frontend institucional", "https://lab.vallegrande.edu.pe/edufront/auth/login")],
    )
    add_project(
        doc,
        "Microservicio de Asistencia",
        "Módulo asignado: asistencia",
        [
            "Servicio orientado a gestionar registros de asistencia dentro del ecosistema EDUNOVA.",
            "Trabajo con arquitectura reactiva, persistencia relacional, contenedores y versionamiento institucional.",
        ],
        [("Repositorio GitLab", "https://gitlab.com/vallegrande/as241s5_prs5/vg-ms-attendance")],
    )
    add_project(
        doc,
        "Dashboard BI - Perfil del Estudiante",
        "Proyecto académico colaborativo",
        [
            "Análisis de datos para visualizar indicadores académicos y de preparación laboral.",
            "Uso de Python, Pandas, Looker Studio y PostgreSQL en un entorno de trabajo colaborativo.",
        ],
        [
            ("Backend", "https://github.com/salvador-quispe/AS241_HF_T5.git"),
            ("Frontend", "https://github.com/salvador-quispe/AS241_HF_T5_FE.git"),
        ],
    )
    add_project(
        doc,
        "Kubernetes & Docker Labs",
        "Prácticas cloud native",
        [
            "Prácticas sobre contenedores, redes, pods, deployments, services y administración de aplicaciones contenerizadas.",
        ],
    )

    add_section_title(doc, "Formación")
    add_bullet(doc, "Estudiante de Análisis de Sistemas.")
    add_bullet(doc, "Proyectos académicos con microservicios, frontend Angular, bases de datos y despliegue.")
    add_bullet(doc, "Idiomas: Español e inglés básico.")

    add_section_title(doc, "Certificaciones y badges")
    certs = [
        "Enterprise Design Thinking Practitioner - IBM / Enterprise Design Thinking.",
        "AWS Academy Graduate - Cloud Foundations - Amazon Web Services Training and Certification.",
        "Introduction to Kubernetes (LFS158) - The Linux Foundation.",
        "Containerization and Virtualization Concepts - DataCamp.",
        "Introducción a Docker - DataCamp.",
    ]
    for cert in certs:
        add_bullet(doc, cert)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("CV generado para portafolio profesional | Victor Raul Torreblanca Franco")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
